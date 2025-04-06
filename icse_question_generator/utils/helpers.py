"""
Utility functions for the ICSE Question Generator.
"""

from typing import List
from question_generator.chemistry import Question
import docx
from docx.shared import Pt
from fpdf import FPDF


def export_to_pdf(questions: List[Question], filename: str) -> None:
    """
    Export questions to a PDF file.
    
    Args:
        questions: List of Question objects
        filename: Output PDF filename
    """
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    
    for i, question in enumerate(questions, 1):
        pdf.cell(200, 10, txt=f"Q{i}. {question.question_text}", ln=True)
        
        if question.options:
            for option in question.options:
                pdf.cell(200, 10, txt=f"   {option}", ln=True)
        
        pdf.ln(5)
    
    pdf.output(filename)


def export_to_docx(questions: List[Question], filename: str) -> None:
    """
    Export questions to a DOCX file.
    
    Args:
        questions: List of Question objects
        filename: Output DOCX filename
    """
    doc = docx.Document()
    
    for i, question in enumerate(questions, 1):
        p = doc.add_paragraph()
        p.add_run(f"Q{i}. {question.question_text}").bold = True
        
        if question.options:
            for option in question.options:
                doc.add_paragraph(f"   {option}", style='List Bullet')
        
        doc.add_paragraph()
    
    doc.save(filename) 